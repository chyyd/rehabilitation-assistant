"""
文档解析器 - 支持PDF、EPUB、Word、TXT
"""
from pathlib import Path
from typing import Optional
import os


class DocumentParser:
    """统一文档解析接口"""

    def parse(self, file_path: str) -> str:
        """自动识别格式并解析为纯文本

        Args:
            file_path: 文件路径

        Returns:
            解析后的文本内容
        """
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return self._parse_word(file_path)
        elif ext == '.epub':
            return self._parse_epub(file_path)
        elif ext == '.txt':
            return self._parse_txt(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def _parse_pdf(self, file_path: str) -> str:
        """解析PDF文件"""
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text

        except ImportError:
            raise ImportError("请安装pypdf库: pip install pypdf")
        except Exception as e:
            raise Exception(f"PDF解析失败: {str(e)}")

    def _parse_word(self, file_path: str) -> str:
        """解析Word文档"""
        try:
            from docx import Document

            doc = Document(file_path)
            text = ""

            # 提取段落文本
            for para in doc.paragraphs:
                text += para.text + "\n"

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    text += " | ".join(row_data) + "\n"

            return text

        except ImportError:
            raise ImportError("请安装python-docx库: pip install python-docx")
        except Exception as e:
            raise Exception(f"Word文档解析失败: {str(e)}")

    def _parse_epub(self, file_path: str) -> str:
        """解析EPUB电子书"""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup

            book = epub.read_epub(file_path)
            text = ""

            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text += soup.get_text() + "\n"

            return text

        except ImportError:
            raise ImportError("请安装ebooklib和beautifulsoup4库")
        except Exception as e:
            raise Exception(f"EPUB解析失败: {str(e)}")

    def _parse_txt(self, file_path: str) -> str:
        """解析纯文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
